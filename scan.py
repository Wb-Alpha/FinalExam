import os
import docx


# 以列表形式返回查重资料库下的所有文件名
def readFilesName(path):
    files_path = path
    name = os.listdir(files_path)
    return name


# 读取文件内容
def readFileContain(path, para):
    content_list = []
    all_files_name = os.listdir(path)  # 读取资源文件夹下的查重资源
    para_list = para.getParameter()

    for file_name in all_files_name:
        # 后期可以添加对.doc文件的支持
        file_path = path + "/" + file_name
        # print(file_path)
        file = docx.Document(file_path)
        content = ""  # 因为docx模块只能一段一段读docx文件中的内容,故增加一个变量用于存储文件内容

        if para_list[0]:    # 扫描文本模式
            for paragraph in file.paragraphs:
                content += paragraph.text
        else:
            tbs = file.tables
            content = tbs[0].rows[7].cells[0].text
            content = content[content.index('<python>') + len('<python>'):content.index('</python>')]

        content_list.append(content)
    return content_list


def loadOriginalText(path):
    file = docx.Document(path)
    content = ""

    for paragraph in file.paragraphs:
        content += paragraph.text

    return content

# readFileContain("D:/网课/python/测试数据")